#!/usr/bin/env python3
"""
Scraping gateway: Scrapling (stealth fetch — bypasses Cloudflare/anti-bot, renders JS) +
ScrapeGraph-ai (LLM extraction). Runs as a separate service so the generated apps stay thin
(no chromium in-app). env: SCRAPER_PORT, OLLAMA_URL, SCRAPER_LLM.

⚠️ Pass session cookies for authenticated/paywalled targets (LinkedIn li_at, etc.). Gray-area:
respect rate limits, rotate proxies, persist tokens — risk of bans (see the OSS research).
"""
import os
from typing import Optional
from fastapi import FastAPI
from pydantic import BaseModel

app = FastAPI()


class FetchReq(BaseModel):
    url: str
    stealth: bool = True
    cookies: Optional[str] = None


@app.post("/fetch")
def fetch(r: FetchReq):
    try:
        from scrapling.fetchers import StealthyFetcher, Fetcher
        if r.stealth:
            page = StealthyFetcher.fetch(r.url, headless=True, network_idle=True)
        else:
            page = Fetcher.get(r.url, stealthy_headers=True)
        html = getattr(page, "html_content", "") or ""
        try:
            text = page.get_all_text()
        except Exception:
            text = ""
        return {"status": getattr(page, "status", 200), "html": html[:500000], "text": (text or "")[:200000]}
    except Exception as e:
        return {"status": 0, "html": "", "text": "", "error": str(e)}


class ExtractReq(BaseModel):
    url: str
    prompt: str
    cookies: Optional[str] = None


@app.post("/extract")
def extract(r: ExtractReq):
    try:
        from scrapegraphai.graphs import SmartScraperGraph
        cfg = {
            "llm": {"model": os.environ.get("SCRAPER_LLM", "ollama/qwen2.5-coder:32b"), "base_url": os.environ.get("OLLAMA_URL", "http://localhost:11434"), "temperature": 0},
            "verbose": False, "headless": True,
        }
        return {"data": SmartScraperGraph(prompt=r.prompt, source=r.url, config=cfg).run()}
    except Exception as e:
        return {"error": str(e)}


class PdfReq(BaseModel):
    html: str
    landscape: bool = False
    format: str = "A4"


@app.post("/pdf")
def pdf(r: PdfReq):
    """HTML → PDF via the same Playwright/chromium (for the pdf module)."""
    try:
        from playwright.sync_api import sync_playwright
        from fastapi import Response
        with sync_playwright() as p:
            b = p.chromium.launch()
            page = b.new_page()
            page.set_content(r.html, wait_until="networkidle")
            data = page.pdf(format=r.format, landscape=r.landscape, print_background=True)
            b.close()
        return Response(content=data, media_type="application/pdf")
    except Exception as e:
        return {"error": str(e)}


class ImageReq(BaseModel):
    url: Optional[str] = None
    data_b64: Optional[str] = None
    width: Optional[int] = None
    height: Optional[int] = None
    quality: int = 82
    fmt: str = "webp"


@app.post("/image")
def image(r: ImageReq):
    """Resize/optimize an image (for the media module). Returns base64 of the result."""
    try:
        import base64, io, urllib.request
        from PIL import Image
        if r.url:
            raw = urllib.request.urlopen(r.url, timeout=20).read()
        else:
            raw = base64.b64decode(r.data_b64 or "")
        im = Image.open(io.BytesIO(raw))
        if r.width or r.height:
            im.thumbnail((r.width or im.width, r.height or im.height))
        out = io.BytesIO()
        im.save(out, format=r.fmt.upper(), quality=r.quality)
        return {"data_b64": base64.b64encode(out.getvalue()).decode(), "format": r.fmt, "width": im.width, "height": im.height}
    except Exception as e:
        return {"error": str(e)}


class OcrReq(BaseModel):
    url: Optional[str] = None
    data_b64: Optional[str] = None
    lang: str = "eng+spa"


@app.post("/ocr")
def ocr(r: OcrReq):
    """Image/scan → text (for the ocr module). Needs tesseract + pytesseract installed."""
    try:
        import base64, io, urllib.request
        import pytesseract
        from PIL import Image
        if r.url:
            raw = urllib.request.urlopen(r.url, timeout=20).read()
        else:
            raw = base64.b64decode(r.data_b64 or "")
        text = pytesseract.image_to_string(Image.open(io.BytesIO(raw)), lang=r.lang)
        return {"text": text.strip()}
    except Exception as e:
        return {"error": str(e)}


class ParseReq(BaseModel):
    url: Optional[str] = None
    data_b64: Optional[str] = None
    filename: str = "doc"


@app.post("/parse")
def parse(r: ParseReq):
    """Any document (PDF/docx/pptx/xlsx/html/image) → clean Markdown (microsoft/markitdown).
    The bridge from files to RAG/LLM. For the docparse module."""
    try:
        from markitdown import MarkItDown
        md = MarkItDown()
        if r.url:
            res = md.convert(r.url)
        else:
            import base64, tempfile, os
            raw = base64.b64decode(r.data_b64 or "")
            ext = os.path.splitext(r.filename)[1] or ".bin"
            with tempfile.NamedTemporaryFile(suffix=ext, delete=False) as f:
                f.write(raw)
                path = f.name
            res = md.convert(path)
        return {"markdown": res.text_content}
    except Exception as e:
        return {"error": str(e)}


class DocgenReq(BaseModel):
    kind: str  # docx | xlsx
    title: Optional[str] = None
    paragraphs: Optional[list] = None      # for docx
    rows: Optional[list] = None            # for xlsx (list of lists)


@app.post("/docgen")
def docgen(r: DocgenReq):
    """Generate Office documents (inspired by anthropics/skills docx/xlsx). Returns base64."""
    try:
        import base64, io
        if r.kind == "xlsx":
            from openpyxl import Workbook
            wb = Workbook(); ws = wb.active
            if r.title:
                ws.title = r.title[:31]
            for row in (r.rows or []):
                ws.append(row)
            out = io.BytesIO(); wb.save(out)
            return {"data_b64": base64.b64encode(out.getvalue()).decode(), "ext": "xlsx"}
        # default docx
        from docx import Document
        doc = Document()
        if r.title:
            doc.add_heading(r.title, 0)
        for p in (r.paragraphs or []):
            doc.add_paragraph(p)
        out = io.BytesIO(); doc.save(out)
        return {"data_b64": base64.b64encode(out.getvalue()).decode(), "ext": "docx"}
    except Exception as e:
        return {"error": str(e)}


if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=int(os.environ.get("SCRAPER_PORT", "8200")))
